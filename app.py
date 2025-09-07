# -*- coding: utf-8 -*-
"""
Streamlit – Formulário para gerar Relatório Técnico (LOGO + Local/Drive/GitHub + Código automático)

Recursos:
- Form com campos principais (metadados, escopo, métodos, resultados, conclusões).
- Campos dinâmicos para autores, referências e anexos.
- Upload de LOGO (PNG/JPG) e controle de largura em cm.
- Preview em Markdown ao vivo.
- Exporta PDF (ReportLab) e DOCX (python-docx) com LOGO.
- Salva/Carrega rascunho em JSON (local) + opção de enviar JSON/MD/PDF/DOCX ao Google Drive e GitHub.
- Gerador automático de código (ex.: MavipeRTEC001).

Como rodar:
1) pip install streamlit reportlab python-docx pydantic pillow google-api-python-client google-auth google-auth-httplib2 requests
2) streamlit run app.py
"""

import io
import os
import base64
import json
import datetime as dt
from pathlib import Path
from typing import List, Optional, Tuple

import requests
import streamlit as st
from pydantic import BaseModel, Field

# ===================== Modelos de Dados =====================
class Autor(BaseModel):
    nome: str = Field("", description="Nome do autor")
    cargo: str = ""
    email: str = ""

class Referencia(BaseModel):
    referencia: str = ""

class Anexo(BaseModel):
    titulo: str = ""
    descricao: str = ""
    link: str = ""  # URL para arquivo/imagem externa

# ===================== Código do relatório: gerador =====================
COUNTER_FILE_DEFAULT = "counter.json"

def next_report_code(prefix: str = "MavipeRTEC", draft_dir: Optional[str] = None) -> str:
    pdir = Path(draft_dir) if draft_dir else Path.cwd() / "drafts"
    pdir.mkdir(parents=True, exist_ok=True)
    cfile = pdir / COUNTER_FILE_DEFAULT
    counter = 0
    if cfile.exists():
        try:
            counter = json.loads(cfile.read_text(encoding="utf-8")).get("counter", 0)
        except Exception:
            counter = 0
    counter += 1
    cfile.write_text(json.dumps({"counter": counter}, ensure_ascii=False, indent=2), encoding="utf-8")
    return f"{prefix}{counter:03d}"

class Relatorio(BaseModel):
    # Metadados
    titulo: str = "Relatório Técnico"
    cliente: str = ""
    projeto: str = ""
    codigo: str = ""  # ex: MavipeRTEC001
    data: str = dt.date.today().strftime("%Y-%m-%d")
    versao: str = "1.0"

    # Autores e aprovadores
    autores: List[Autor] = Field(default_factory=lambda: [Autor(nome="", cargo="", email="")])
    aprovador: str = ""

    # Corpo
    resumo_exec: str = ""
    escopo: str = ""
    dados_fontes: str = ""  # fontes de dados utilizados
    metodologia: str = ""
    resultados: str = ""
    discussoes: str = ""
    conclusoes: str = ""
    recomendacoes: str = ""

    # Outros
    referencias: List[Referencia] = Field(default_factory=list)
    anexos: List[Anexo] = Field(default_factory=list)
    observacoes: str = ""

# ===================== Utilidades =====================
def to_markdown(r: Relatorio) -> str:
    """Gera um markdown organizado do relatório."""
    autores_md = "\n".join([f"- {a.nome} ({a.cargo}) <{a.email}>" for a in r.autores if a.nome.strip()])
    refs_md = "\n".join([f"- {x.referencia}" for x in r.referencias if x.referencia.strip()])
    anexos_md = "\n".join([
        f"- **{a.titulo}** – {a.descricao} {(f'({a.link})' if a.link else '')}"
        for a in r.anexos if a.titulo.strip()
    ])

    parts = [
        f"# {r.titulo}",
        f"**Cliente:** {r.cliente}  ",
        f"**Projeto:** {r.projeto}  ",
        f"**Código:** {r.codigo}  ",
        f"**Data:** {r.data}  ",
        f"**Versão:** {r.versao}",
        "\n---\n",
        "## Autores",
        autores_md or "(preencher)",
        f"\n**Aprovador:** {r.aprovador or '(preencher)'}\n",
        "\n## Resumo Executivo\n" + (r.resumo_exec or "(preencher)"),
        "\n## Escopo\n" + (r.escopo or "(preencher)"),
        "\n## Dados & Fontes\n" + (r.dados_fontes or "(preencher)"),
        "\n## Metodologia\n" + (r.metodologia or "(preencher)"),
        "\n## Resultados\n" + (r.resultados or "(preencher)"),
        "\n## Discussões\n" + (r.discussoes or "(preencher)"),
        "\n## Conclusões\n" + (r.conclusoes or "(preencher)"),
        "\n## Recomendações\n" + (r.recomendacoes or "(preencher)"),
        "\n## Referências\n" + (refs_md or "(preencher)"),
        "\n## Anexos\n" + (anexos_md or "(preencher)"),
        "\n## Observações\n" + (r.observacoes or ""),
    ]
    return "\n".join(parts)

# ===================== Logo helpers =====================
def get_logo_dims_cm(logo_bytes: bytes, width_cm: float) -> Tuple[float, float]:
    """Calcula altura proporcional em cm, mantendo aspecto."""
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(logo_bytes))
    w, h = img.size
    if w == 0 or h == 0:
        return width_cm, width_cm * 0.5
    ratio = h / w
    return width_cm, width_cm * ratio

# ===================== Exportadores =====================
def build_pdf(r: Relatorio, logo_bytes: Optional[bytes], logo_width_cm: float) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.units import cm

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm
    )
    styles = getSampleStyleSheet()
    story = []

    def p(text: str, style: str = "BodyText"):
        story.append(Paragraph(text.replace("\n", "<br/>"), styles[style]))
        story.append(Spacer(1, 0.3*cm))

    # LOGO (se houver)
    if logo_bytes:
        w_cm, h_cm = get_logo_dims_cm(logo_bytes, logo_width_cm)
        story.append(RLImage(io.BytesIO(logo_bytes), width=w_cm*cm, height=h_cm*cm))
        story.append(Spacer(1, 0.4*cm))

    # Cabeçalho
    p(f"<b>{r.titulo}</b>", "Title")
    p(
        f"Cliente: <b>{r.cliente or '-'}"
        + "</b><br/>"
        + f"Projeto: <b>{r.projeto or '-'}"
        + "</b><br/>"
        + f"Código: <b>{r.codigo or '-'}"
        + "</b><br/>"
        + f"Data: <b>{r.data or '-'}"
        + "</b><br/>"
        + f"Versão: <b>{r.versao or '-'}"
        + "</b>"
    )

    # Autores
    autores = "<br/>".join([f"- {a.nome} ({a.cargo}) &lt;{a.email}&gt;" for a in r.autores if a.nome.strip()]) or "(preencher)"
    p(f"<b>Autores</b><br/>{autores}")
    p(f"<b>Aprovador</b><br/>{r.aprovador or '(preencher)'}")

    # Seções
    def sec(title: str, text: str):
        p(f"<b>{title}</b>")
        p(text or "(preencher)")

    sec("Resumo Executivo", r.resumo_exec)
    sec("Escopo", r.escopo)
    sec("Dados & Fontes", r.dados_fontes)
    sec("Metodologia", r.metodologia)
    sec("Resultados", r.resultados)
    sec("Discussões", r.discussoes)
    sec("Conclusões", r.conclusoes)
    sec("Recomendações", r.recomendacoes)

    # Referências e Anexos
    refs = "<br/>".join([f"- {x.referencia}" for x in r.referencias if x.referencia.strip()]) or "(preencher)"
    p(f"<b>Referências</b><br/>{refs}")

    anexos = "<br/>".join(
        [f"- <b>{a.titulo}</b> – {a.descricao} {(f'({a.link})' if a.link else '')}" for a in r.anexos if a.titulo.strip()]
    ) or "(preencher)"
    p(f"<b>Anexos</b><br/>{anexos}")

    if r.observacoes:
        sec("Observações", r.observacoes)

    doc.build(story)
    return buffer.getvalue()

def build_docx(r: Relatorio, logo_bytes: Optional[bytes], logo_width_cm: float) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cabeçalho com LOGO (se houver)
    if logo_bytes:
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(logo_bytes), width=Cm(logo_width_cm))

    doc.add_heading(r.titulo or "Relatório Técnico", level=0)

    meta = doc.add_paragraph()
    meta.add_run("Cliente: ").bold = True; meta.add_run(r.cliente or "-")
    meta.add_run("\nProjeto: ").bold = True; meta.add_run(r.projeto or "-")
    meta.add_run("\nCódigo: ").bold = True; meta.add_run(r.codigo or "-")
    meta.add_run("\nData: ").bold = True; meta.add_run(r.data or "-")
    meta.add_run("\nVersão: ").bold = True; meta.add_run(r.versao or "-")

    doc.add_heading("Autores", level=1)
    for a in r.autores:
        if a.nome.strip():
            doc.add_paragraph(f"- {a.nome} ({a.cargo}) <{a.email}>")
    doc.add_paragraph(f"Aprovador: {r.aprovador or '(preencher)'}")

    def sec(title: str, text: str):
        doc.add_heading(title, level=1)
        doc.add_paragraph(text or "(preencher)")

    sec("Resumo Executivo", r.resumo_exec)
    sec("Escopo", r.escopo)
    sec("Dados & Fontes", r.dados_fontes)
    sec("Metodologia", r.metodologia)
    sec("Resultados", r.resultados)
    sec("Discussões", r.discussoes)
    sec("Conclusões", r.conclusoes)
    sec("Recomendações", r.recomendacoes)

    doc.add_heading("Referências", level=1)
    if r.referencias:
        for x in r.referencias:
            if x.referencia.strip():
                doc.add_paragraph(f"- {x.referencia}")
    else:
        doc.add_paragraph("(preencher)")

    doc.add_heading("Anexos", level=1)
    if r.anexos:
        for a in r.anexos:
            if a.titulo.strip():
                line = f"- {a.titulo} – {a.descricao}"
                if a.link:
                    line += f" ({a.link})"
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("(preencher)")

    if r.observacoes:
        sec("Observações", r.observacoes)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ===================== Google Drive (opcional) =====================
# Em st.secrets:
# [gcp_service_account]
# type = "service_account"
# project_id = "..."
# private_key_id = "..."
# private_key = "-----BEGIN PRIVATE KEY-----\n...\n-----END PRIVATE KEY-----\n"
# client_email = "...@...gserviceaccount.com"
# client_id = "..."
# token_uri = "https://oauth2.googleapis.com/token"
# [drive]
# folder_id = "<ID da pasta de destino>"

def get_drive_service():
    try:
        from google.oauth2.service_account import Credentials
        from googleapiclient.discovery import build
        scopes = ["https://www.googleapis.com/auth/drive.file"]
        sa_info = dict(st.secrets["gcp_service_account"])  # type: ignore
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        service = build("drive", "v3", credentials=creds)
        return service
    except Exception as e:
        st.error(f"Drive não configurado: {e}")
        return None

def drive_upload_bytes(service, folder_id: str, filename: str, data: bytes, mime: str) -> str:
    from googleapiclient.http import MediaIoBaseUpload
    import io as _io
    file_metadata = {"name": filename, "parents": [folder_id]}
    media = MediaIoBaseUpload(_io.BytesIO(data), mimetype=mime, resumable=False)
    f = service.files().create(body=file_metadata, media_body=media, fields="id, webViewLink").execute()
    return f.get("webViewLink", "")

# ===================== GitHub (opcional) =====================
# Em st.secrets:
# [github]
# token = "ghp_..."
# repo = "usuario/repositorio"
# branch = "main"
# base_path = "reports"  # opcional

def github_upload_bytes(filename: str, data: bytes, message: str) -> str:
    try:
        gh = st.secrets.get("github", {})
        token = gh.get("token"); repo = gh.get("repo"); branch = gh.get("branch", "main"); base_path = gh.get("base_path", "reports")
        if not token or not repo:
            raise RuntimeError("Token/Repo não configurados em st.secrets['github']")
        url = f"https://api.github.com/repos/{repo}/contents/{base_path}/{filename}"
        payload = {"message": message, "content": base64.b64encode(data).decode("utf-8"), "branch": branch}
        r = requests.put(url, json=payload, headers={"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}, timeout=30)
        if r.status_code not in (200, 201):
            raise RuntimeError(f"GitHub API: {r.status_code} – {r.text}")
        j = r.json()
        return j.get("content", {}).get("html_url", "") or j.get("content", {}).get("path", "")
    except Exception as e:
        st.error(f"GitHub upload falhou: {e}")
        return ""

# ===================== UI =====================
st.set_page_config(page_title="Formulário – Relatório Técnico", page_icon="📝", layout="wide")

st.markdown(
    """
    <style>
    .block-container{padding-top:1rem;padding-bottom:2rem;}
    .stTextArea textarea{min-height:120px}
    .small{opacity:.75;font-size:.9rem}
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("📝 Formulário rápido para Relatório Técnico")
st.caption("Preencha os campos e exporte em PDF/DOCX/Markdown. Também é possível salvar rascunho local ou enviar ao Drive/GitHub.")

# -------- Sidebar --------
with st.sidebar:
    st.header("Ações")
    st.write("Logo (opcional), rascunhos e salvamento local/Drive/GitHub.")

    # Estado inicial
    if "rel" not in st.session_state:
        st.session_state.rel = Relatorio()
    rel: Relatorio = st.session_state.rel

    # Gerador de código
    code_prefix = st.text_input("Prefixo do código", value=st.session_state.get("code_prefix", "MavipeRTEC"))
    st.session_state.code_prefix = code_prefix
    if st.button("🔢 Gerar código", use_container_width=True):
        new_code = next_report_code(prefix=code_prefix, draft_dir=st.session_state.get("draft_dir"))
        rel.codigo = new_code
        st.session_state.rel = rel
        st.success(f"Código gerado: {new_code}")

    st.markdown("---")

    # Destino (info apenas visual; o salvamento local funciona independente)
    st.radio("Destino padrão para rascunhos", options=["Local", "Google Drive", "GitHub"], index=0, horizontal=False, key="dest")

    st.markdown("---")

    # Salvamento LOCAL
    st.subheader("Salvar local")
    default_dir = st.session_state.get("draft_dir", str(Path.cwd() / "drafts"))
    draft_dir = st.text_input("Pasta de rascunhos", value=default_dir, help="Será criada se não existir.")
    st.session_state.draft_dir = draft_dir
    autosave = st.checkbox("Salvar local ao atualizar prévia", value=st.session_state.get("autosave", True))
    st.session_state.autosave = autosave

    # Listar rascunhos
    try:
        p = Path(draft_dir)
        p.mkdir(parents=True, exist_ok=True)
        available = sorted([f.name for f in p.glob("*.json")])
    except Exception as _e:
        available = []
        st.error(f"Não foi possível acessar/criar a pasta: {_e}")
    load_choice = st.selectbox("Carregar rascunho local", options=["(nenhum)"] + available)
    if load_choice != "(nenhum)":
        try:
            data = json.loads((Path(draft_dir) / load_choice).read_text(encoding="utf-8"))
            st.session_state.rel = Relatorio(**data)
            st.success(f"Rascunho carregado: {load_choice}")
            rel = st.session_state.rel
        except Exception as e:
            st.error(f"Erro ao carregar rascunho: {e}")

    if st.button("💾 Salvar local agora", use_container_width=True):
        try:
            p = Path(draft_dir); p.mkdir(parents=True, exist_ok=True)
            fname = f"{(st.session_state.rel.codigo or 'relatorio').replace(' ', '_')}.json"
            (p / fname).write_text(json.dumps(st.session_state.rel.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
            st.success(f"Salvo em {p / fname}")
        except Exception as e:
            st.error(f"Falha ao salvar: {e}")

    st.markdown("---")

    # Logo
    st.subheader("Logo")
    logo_file = st.file_uploader("Logo (PNG/JPG)", type=["png", "jpg", "jpeg"])
    logo_width_cm = st.number_input("Largura do logo (cm)", min_value=1.0, max_value=12.0, value=3.5, step=0.5)
    if "logo_bytes" not in st.session_state:
        st.session_state.logo_bytes = None
    if logo_file is not None:
        st.session_state.logo_bytes = logo_file.read()
        st.success("Logo carregado!")

    st.markdown("---")

    # Upload/download de rascunho (JSON)
    st.write("Upload/Download de rascunhos (JSON)")
    empty_model = Relatorio().model_dump()
    st.download_button(
        label="⬇️ Baixar modelo JSON",
        data=json.dumps(empty_model, ensure_ascii=False, indent=2).encode("utf-8"),
        file_name="relatorio_modelo.json",
        mime="application/json",
        use_container_width=True,
    )

    up = st.file_uploader("Importar rascunho JSON", type=["json"], help="Carregue um arquivo salvo anteriormente")
    if up is not None:
        try:
            data = json.load(up)
            imported = Relatorio(**data)
            st.session_state.rel = imported
            st.success("Rascunho importado!")
        except Exception as e:
            st.error(f"Erro ao importar: {e}")

# -------- Form principal --------
rel: Relatorio = st.session_state.rel
with st.form("form-relatorio"):
    st.subheader("Metadados")
    c1, c2, c3, c4 = st.columns([2, 2, 1, 1])
    rel.titulo = c1.text_input("Título", value=rel.titulo)
    rel.cliente = c2.text_input("Cliente", value=rel.cliente)
    rel.projeto = c3.text_input("Projeto", value=rel.projeto)
    rel.codigo = c4.text_input("Código", value=rel.codigo)

    d1, d2 = st.columns(2)
    rel.data = d1.date_input("Data", value=dt.date.fromisoformat(rel.data) if rel.data else dt.date.today()).isoformat()
    rel.versao = d2.text_input("Versão", value=rel.versao)

    st.markdown("---")

    st.subheader("Autores & Aprovação")
    n_aut = st.number_input("Nº de autores", min_value=1, max_value=10, value=max(1, len(rel.autores)), step=1)
    while len(rel.autores) < n_aut:
        rel.autores.append(Autor(nome="", cargo="", email=""))
    while len(rel.autores) > n_aut:
        rel.autores.pop()

    for i, a in enumerate(rel.autores):
        c1, c2, c3 = st.columns([2, 2, 2])
        a.nome = c1.text_input(f"Autor {i+1} – Nome", value=a.nome)
        a.cargo = c2.text_input(f"Autor {i+1} – Cargo", value=a.cargo)
        a.email = c3.text_input(f"Autor {i+1} – E-mail", value=a.email)

    rel.aprovador = st.text_input("Aprovador", value=rel.aprovador)

    st.markdown("---")

    st.subheader("Conteúdo")
    rel.resumo_exec   = st.text_area("Resumo Executivo", value=rel.resumo_exec, placeholder="Contexto, objetivo, principais achados, 5–8 linhas.")
    rel.escopo        = st.text_area("Escopo", value=rel.escopo)
    rel.dados_fontes  = st.text_area("Dados & Fontes", value=rel.dados_fontes, placeholder="Ex.: Sentinel-1 (S1), AIS, GNSS, bases internas…")
    rel.metodologia   = st.text_area("Metodologia", value=rel.metodologia, placeholder="Ex.: InSAR PS, CFAR 100m tiles, limiares…")
    rel.resultados    = st.text_area("Resultados", value=rel.resultados)
    rel.discussoes    = st.text_area("Discussões", value=rel.discussoes)
    rel.conclusoes    = st.text_area("Conclusões", value=rel.conclusoes)
    rel.recomendacoes = st.text_area("Recomendações", value=rel.recomendacoes)

    st.markdown("---")

    st.subheader("Referências")
    n_ref = st.number_input("Nº de referências", min_value=0, max_value=50, value=len(rel.referencias) if rel.referencias else 0, step=1)
    while len(rel.referencias) < n_ref:
        rel.referencias.append(Referencia(referencia=""))
    while len(rel.referencias) > n_ref:
        rel.referencias.pop()
    for i, rref in enumerate(rel.referencias):
        rref.referencia = st.text_input(f"Ref. {i+1}", value=rref.referencia)

    st.subheader("Anexos")
    n_anx = st.number_input("Nº de anexos", min_value=0, max_value=30, value=len(rel.anexos) if rel.anexos else 0, step=1)
    while len(rel.anexos) < n_anx:
        rel.anexos.append(Anexo(titulo="", descricao="", link=""))
    while len(rel.anexos) > n_anx:
        rel.anexos.pop()
    for i, anx in enumerate(rel.anexos):
        c1, c2 = st.columns([2, 3])
        anx.titulo = c1.text_input(f"Anexo {i+1} – Título", value=anx.titulo)
        anx.descricao = c2.text_input(f"Anexo {i+1} – Descrição", value=anx.descricao)
        anx.link = st.text_input(f"Anexo {i+1} – Link (opcional)", value=anx.link)

    rel.observacoes = st.text_area("Observações (opcional)", value=rel.observacoes)

    submitted = st.form_submit_button("Atualizar prévia")

    # Autosave local
    if submitted and st.session_state.get("autosave"):
        try:
            p = Path(st.session_state.get("draft_dir", str(Path.cwd() / "drafts")))
            p.mkdir(parents=True, exist_ok=True)
            fname = f"{(rel.codigo or 'relatorio').replace(' ', '_')}.json"
            (p / fname).write_text(json.dumps(rel.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
            st.toast("Rascunho salvo localmente", icon="💾")
        except Exception as e:
            st.warning(f"Não foi possível salvar automaticamente: {e}")

# Atualiza sessão
st.session_state.rel = rel

# Prévia em Markdown
st.subheader("Prévia (Markdown)")
st.markdown("<span class='small'>Copie e cole onde quiser ou exporte pelos botões abaixo.</span>", unsafe_allow_html=True)
st.code(to_markdown(rel), language="markdown")

# Botões de exportação
colA, colB, colC = st.columns([1, 1, 1])

md_bytes = to_markdown(rel).encode("utf-8")
colA.download_button("⬇️ Baixar .md", data=md_bytes, file_name=f"{(rel.codigo or 'relatorio').replace(' ', '_')}.md", mime="text/markdown", use_container_width=True)

try:
    pdf_bytes = build_pdf(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
except Exception:
    # Em alguns ambientes, o number_input não persistiu a largura no session_state
    pdf_bytes = build_pdf(rel, st.session_state.get("logo_bytes"), 3.5)

try:
    colB.download_button("⬇️ Baixar PDF", data=pdf_bytes, file_name=f"{(rel.codigo or 'relatorio')}.pdf", mime="application/pdf", use_container_width=True)
except Exception as e:
    colB.error(f"Erro PDF: {e}")

try:
    docx_bytes = build_docx(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
    colC.download_button("⬇️ Baixar DOCX", data=docx_bytes, file_name=f"{(rel.codigo or 'relatorio')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
except Exception as e:
    colC.error(f"Erro DOCX: {e}")

st.markdown("---")

# ====== Salvar no Google Drive (opcional) ======
st.subheader("Salvar no Google Drive (opcional)")
if st.secrets.get("gcp_service_account") and st.secrets.get("drive", {}).get("folder_id"):
    if st.button("📤 Enviar rascunho (.json) ao Drive", use_container_width=True):
        try:
            svc = get_drive_service()
            folder_id = st.secrets["drive"]["folder_id"]
            base = (rel.codigo or "relatorio").replace(" ", "_")
            json_bytes = json.dumps(rel.model_dump(), ensure_ascii=False, indent=2).encode("utf-8")
            url_json = drive_upload_bytes(svc, folder_id, f"{base}.json", json_bytes, "application/json")
            st.success("Rascunho enviado!")
            st.markdown(f"- **JSON:** [abrir]({url_json})")
        except Exception as e:
            st.error(f"Falha ao enviar ao Drive: {e}")
    if st.button("📤 Enviar exportações (MD/PDF/DOCX) ao Drive", use_container_width=True):
        try:
            md_bytes_up = to_markdown(rel).encode("utf-8")
            pdf_bytes_up = build_pdf(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
            docx_bytes_up = build_docx(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
            svc = get_drive_service()
            folder_id = st.secrets["drive"]["folder_id"]
            base = (rel.codigo or "relatorio").replace(" ", "_")
            links = []
            links.append(("Markdown", drive_upload_bytes(svc, folder_id, f"{base}.md", md_bytes_up, "text/markdown")))
            links.append(("PDF", drive_upload_bytes(svc, folder_id, f"{base}.pdf", pdf_bytes_up, "application/pdf")))
            links.append(("DOCX", drive_upload_bytes(svc, folder_id, f"{base}.docx", docx_bytes_up, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")))
            st.success("Exportações enviadas para o Drive:")
            for label, url in links:
                st.markdown(f"- **{label}:** [abrir]({url})")
        except Exception as e:
            st.error(f"Falha ao enviar ao Drive: {e}")
else:
    st.caption("Para ativar o Drive, configure Service Account e folder_id em st.secrets.")

st.markdown("---")

# ====== GitHub (opcional) ======
st.subheader("Salvar no GitHub (opcional)")
if st.secrets.get("github"):
    if st.button("🐙 Enviar rascunho (.json) ao GitHub", use_container_width=True):
        try:
            base = (rel.codigo or "relatorio").replace(" ", "_")
            json_bytes = json.dumps(rel.model_dump(), ensure_ascii=False, indent=2).encode("utf-8")
            url_json = github_upload_bytes(f"{base}.json", json_bytes, f"chore: rascunho {base}")
            st.success("Rascunho enviado ao GitHub!")
            if url_json:
                st.markdown(f"- **JSON:** {url_json}")
        except Exception as e:
            st.error(f"Falha GitHub: {e}")
    if st.button("🐙 Enviar exportações (MD/PDF/DOCX) ao GitHub", use_container_width=True):
        try:
            base = (rel.codigo or "relatorio").replace(" ", "_")
            md_bytes_up = to_markdown(rel).encode("utf-8")
            pdf_bytes_up = build_pdf(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
            docx_bytes_up = build_docx(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
            u1 = github_upload_bytes(f"{base}.md", md_bytes_up, f"feat: relatório {base} (MD)")
            u2 = github_upload_bytes(f"{base}.pdf", pdf_bytes_up, f"feat: relatório {base} (PDF)")
            u3 = github_upload_bytes(f"{base}.docx", docx_bytes_up, f"feat: relatório {base} (DOCX)")
            st.success("Exportações enviadas ao GitHub!")
            for label, url in [("Markdown", u1), ("PDF", u2), ("DOCX", u3)]:
                if url:
                    st.markdown(f"- **{label}:** {url}")
        except Exception as e:
            st.error(f"Falha GitHub: {e}")
else:
    st.caption("Para ativar o GitHub, configure token/repo/branch/base_path em st.secrets.")

st.divider()

# Download do rascunho atual
st.download_button(
    "💾 Baixar rascunho (JSON)",
    data=json.dumps(rel.model_dump(), ensure_ascii=False, indent=2).encode("utf-8"),
    file_name=f"{(rel.codigo or 'relatorio')}_rascunho.json",
    mime="application/json",
    use_container_width=True,
)

st.caption("Dica: personalize as seções direto no código para criar templates por cliente (OGMP, InSAR, Portos, etc.).")
